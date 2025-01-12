from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
from werkzeug.utils import secure_filename
import os
import csv
import uuid
from hashlib import sha256
import mysql.connector
from mysql.connector import Error
from PyPDF2 import PdfReader
from docx import Document
from google.generativeai import GenerativeModel, configure

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
import json
from datetime import datetime, timedelta

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


app = Flask(__name__)
# AI Configuration
app.secret_key = 'AIzaSyBP45y-sVM_W4T3yacnCRYOr4d6-CgF7A8'


# Custom filter to escape JavaScript strings
@app.template_filter('escapejs')
def escapejs(value):
    if isinstance(value, str):
        return json.dumps(value)[1:-1]  # Escapes JavaScript strings using JSON
    return value


# Configuration
UPLOAD_DIR = "uploads"
REFINED_DIR = "refined"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REFINED_DIR, exist_ok=True)

# MySQL Configuration
MYSQL_HOST = os.getenv('MYSQL_HOST', 'localhost')
MYSQL_USER = os.getenv('MYSQL_USER', 'root')
MYSQL_PASSWORD = os.getenv('MYSQL_PASSWORD', '')
MYSQL_DATABASE = os.getenv('MYSQL_DATABASE', 'resume_tool')

# AI Configuration
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY', 'AIzaSyBP45y-sVM_W4T3yacnCRYOr4d6-CgF7A8')
configure(api_key=GOOGLE_API_KEY)

# Database Connection
def get_db_connection():
    return mysql.connector.connect(
        host=MYSQL_HOST,
        user=MYSQL_USER,
        password=MYSQL_PASSWORD,
        database=MYSQL_DATABASE
    )

# User Authentication
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        try:
            conn = get_db_connection()
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            user = cursor.fetchone()
            
            if user and sha256(password.encode()).hexdigest() == user['password_hash']:
                # Store user_id in session
                session['username'] = username
                session['user_id'] = user['id']  # Assuming 'id' is the user ID column
                flash('Login successful!', 'success')
                return redirect(url_for('dashboard'))
            else:
                flash('Invalid username or password', 'error')
        except Error as e:
            flash(f'Database error: {e}', 'error')
        finally:
            if conn:
                conn.close()
                
    return render_template('login.html')


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        
        try:
            password_hash = sha256(password.encode()).hexdigest()
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO users (username, email, password_hash) VALUES (%s, %s, %s)",
                (username, email, password_hash)
            )
            conn.commit()
            flash('Registration successful! Please login.', 'success')
            return redirect(url_for('login'))
        except Error as e:
            flash(f'Registration error: {e}', 'error')
        finally:
            if conn:
                conn.close()
                
    return render_template('signup.html')

@app.route('/logout')
def logout():
    session.pop('username', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

# Main Routes
@app.route('/')
def index():
    if 'username' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))


# Route to display the dashboard
@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    # Get user ID from session
    user_id = session.get('user_id')

    if not user_id:
        # Handle the case where the user is not logged in
        return redirect(url_for('login'))

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Get resumes for the logged-in user
    cursor.execute("""
        SELECT id 
        FROM resumes 
        WHERE user_id = %s
    """, (user_id,))
    user_resumes = cursor.fetchall()

    # Count the total number of resumes for the user
    total_resumes = len(user_resumes)

    # Get average processing time for the resumes that have been refined
    cursor.execute("""
        SELECT 
            AVG(TIMESTAMPDIFF(SECOND, created_at, NOW())) / 60 AS avg_processing_time
        FROM resumes
        WHERE refined_resume_path IS NOT NULL AND user_id = %s
    """, (user_id,))
    avg_processing_time = cursor.fetchone()['avg_processing_time']
    avg_processing_time = round(avg_processing_time, 2) if avg_processing_time else "N/A"

    # Get the number of active users
    cursor.execute("""
        SELECT COUNT(*) AS active_users 
        FROM users 
        WHERE last_active > NOW() - INTERVAL 1 DAY
    """)
    active_users = cursor.fetchone()['active_users']

    # Get the success rate for resumes that have been refined
    cursor.execute("""
        SELECT 
            (COUNT(refined_resume_path) / COUNT(*)) * 100 AS success_rate
        FROM resumes
        WHERE user_id = %s
    """, (user_id,))
    success_rate = cursor.fetchone()['success_rate']
    success_rate = round(success_rate, 2) if success_rate else 0

    # Get recent activity for the logged-in user
    cursor.execute("""
        SELECT activity_type, timestamp 
        FROM activity_logs 
        WHERE user_id = %s 
        ORDER BY timestamp DESC
        LIMIT 5
    """, (user_id,))
    recent_activity = cursor.fetchall()

    # Engagement calculation logic
    user_engagement = calculate_user_engagement(user_id)
    
    return render_template('dashboard.html', 
                           total_resumes=total_resumes, 
                           avg_processing_time=avg_processing_time,
                           active_users=active_users, 
                           success_rate=success_rate,
                           recent_activity=recent_activity,
                           user_engagement=user_engagement,
                           user_id=user_id)


# Score/Engagement calculation endpoint
@app.route('/calculate_score', methods=['POST'])
def calculate_score():
    data = request.get_json()
    total_resumes = data['total_resumes']
    user_id = data['user_id']

    # Calculate engagement score based on resume count
    score = calculate_user_engagement(user_id, total_resumes)
    
    return jsonify({'score': score})


def calculate_user_engagement(user_id, total_resumes=None):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Get the number of actions the user has taken in the last week
    week_start = datetime.now() - timedelta(weeks=1)
    cursor.execute("""
        SELECT COUNT(*) AS actions_last_week
        FROM activity_logs 
        WHERE user_id = %s AND timestamp > %s
    """, (user_id, week_start))
    actions_last_week = cursor.fetchone()['actions_last_week']
    
    # Example: Give more weight to certain types of actions (e.g., feedback)
    cursor.execute("""
        SELECT COUNT(*) AS feedback_count
        FROM activity_logs
        WHERE user_id = %s AND activity_type = 'feedback' AND timestamp > %s
    """, (user_id, week_start))
    feedback_count = cursor.fetchone()['feedback_count']

    # Calculate engagement score based on actions and feedback
    base_score = actions_last_week * 0.7 + feedback_count * 1.5  # Weight actions and feedback

    # Normalize the score by the total number of resumes (or another metric)
    if total_resumes is not None:
        engagement_score = (base_score / total_resumes) * 100
    else:
        # Default logic in case total_resumes isn't provided
        engagement_score = (base_score / 10) * 100  # Default scaling factor

    # Ensure the score doesn't exceed 100%
    engagement_score = min(engagement_score, 100)

    return round(engagement_score, 2)




def extract_text_from_pdf(pdf_file):
    """
    Extract text from PDF with improved error handling
    """
    try:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                logger.error(f"Error extracting text from page: {e}")
                continue
        
        if not text.strip():
            logger.error("No text extracted from PDF")
            flash("No readable text found in the PDF. Please ensure the PDF contains extractable text.", 'error')
            return None
            
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        flash(f"Error processing PDF: {str(e)}", 'error')
        return None


def extract_text_from_docx(docx_file):
    """
    Extract text from DOCX file with improved error handling
    """
    try:
        doc = Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"

        if not text.strip():
            logger.error("No text extracted from DOCX")
            flash("No readable text found in the DOCX file. Please ensure the DOCX contains extractable text.", 'error')
            return None
        
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        flash(f"Error processing DOCX: {str(e)}", 'error')
        return None


def extract_text_from_file(file):
    """
    Extract text from either PDF or DOCX file based on file type
    """
    if file.filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif file.filename.endswith('.docx'):
        return extract_text_from_docx(file)
    else:
        flash("Unsupported file type. Please upload a PDF or DOCX file.", 'error')
        return None


def refine_resume(job_desc, resume_text):
    """
    Refine resume with improved error handling and validation
    """
    if not job_desc or not resume_text:
        logger.error("Missing job description or resume text")
        flash("Both job description and resume are required", 'error')
        return None, None

    try:
        model = GenerativeModel(model_name="gemini-1.5-pro")
        chat_session = model.start_chat(history=[])
        
        # First, validate the inputs aren't too long
        if len(job_desc) > 10000 or len(resume_text) > 10000:
            flash("Job description or resume text is too long. Please provide shorter text.", 'error')
            return None, None

        prompt = f"""
        You are an expert in ATS optimization and resume tailoring. Your task is to optimize resumes to align with job requirements while maintaining authenticity and highlighting key qualifications.

        Analysis:
        1. Job Requirements: Identify key skills, experiences, and industry-specific terms.
        2. Resume Gap: Compare qualifications with job requirements and identify alignment.

        Optimization:
        1. Prioritize relevant qualifications and achievements.
        2. Reorganize content with a 65-35 ratio of technical/role-specific to transferable skills.
        3. Quantify results with metrics (%, $, time saved).
        4. Optimize for ATS: Use strategic keyword placement, both full terms and acronyms, with a 5-7% keyword density.
        5. Ensure clear formatting, bullet points for experience, and consistent date formats.

        Output:
        Provide ONLY the optimized resume in a clean, professional format. Focus on content that employers will see, using standard section headers (Summary, Experience, Education, Skills, etc.) and clear bullet points.

        Input:
        Job Description: {job_desc}
        Current Resume: {resume_text}

        Note: The focus is solely on creating a polished, ATS-optimized resume.
        """


        response = chat_session.send_message(prompt)
        refined_text = response.text.strip()

        if not refined_text:
            logger.error("Empty response from AI model")
            flash("Error: Unable to generate refined resume", 'error')
            return None, None

        # Get job category
        category_response = chat_session.send_message(
            f"Based on this job description, provide a single job category (1-3 words):\n{job_desc}"
        )
        job_category = category_response.text.strip()

        return refined_text, job_category

    except Exception as e:
        logger.error(f"AI refinement error: {e}")
        flash(f"Error during resume refinement: {str(e)}", 'error')
        return None, None
    

def create_docx(resume_text):
    """
    Create a professional DOCX with improved error handling and formatting
    """
    try:
        doc = Document()

        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(72)  # 1 inch
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)


        # Add content with proper formatting
        lines = resume_text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this is a section header (e.g., Experience, Education)
            if line.isupper() or any(header in line.lower() for header in ['experience', 'education', 'skills', 'summary']):
                if current_section:  # Add spacing between sections
                    doc.add_paragraph()
                current_section = doc.add_paragraph()
                section_run = current_section.add_run(line)
                section_run.font.size = Pt(14)
                section_run.font.bold = True
                section_run.font.color.rgb = RGBColor(0, 51, 102)
                current_section.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                doc.add_paragraph()  # Add spacing
            else:
                # Add content under each section
                para = doc.add_paragraph()
                para_run = para.add_run(line)
                para_run.font.size = Pt(12)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # Save the document to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    except Exception as e:
        logger.error(f"DOCX creation error: {e}")
        flash(f"Error creating document: {str(e)}", 'error')
        return None


def categorize_resume(job_desc):
    prompt = f"""
    Analyze the following job description and provide ONLY a single job category (1-3 words) that best describes the role.
    No additional explanation or context - just the category itself.

    Job Description:
    {job_desc}
    """

    try:
        chat_session = GenerativeModel(model_name="gemini-1.5-pro").start_chat(history=[])
        response = chat_session.send_message(prompt)
        job_category = response.text.strip()

        if not job_category:
            flash("No job category returned by AI", 'error')
            return None

        return job_category
    except Exception as e:
        flash(f"AI Error: {e}", 'error')
        return None
         



def save_refined_resume(username, resume_path, resume_text, job_description, refined_resume_path, job_category):    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Retrieve the user_id based on the username
        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        if user is None:
            flash("User not found", 'error')
            return
        
        user_id = user[0]  # Get the user ID

        if not job_category:
            flash("Job category is missing. Cannot save refined resume.", 'error')
            return
        
        # Insert the refined resume data, including job category
        cursor.execute(
            "INSERT INTO resumes (user_id, resume_path, resume_text, job_description, refined_resume_path, job_category, created_at) VALUES (%s, %s, %s, %s, %s, %s, NOW())",
            (user_id, resume_path, resume_text, job_description, refined_resume_path, job_category)
        )
        conn.commit()
        flash("Refined resume saved successfully!", 'success')
    except Error as e:
        flash(f"Database error while saving the refined resume: {e}", 'error')
    finally:
        if conn:
            resume_id = cursor.lastrowid
            conn.close()
            return resume_id
            

            



@app.route('/new-refinement', methods=['GET', 'POST'])
def new_refinement():
    if 'username' not in session:
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        if 'resume' not in request.files:
            flash('No resume file uploaded', 'error')
            return redirect(request.url)
            
        file = request.files['resume']
        job_desc = request.form['job_description']
        
        if file.filename == '':
            flash('No selected file', 'error')
            return redirect(request.url)
            
        if file:
            try:
                # Process the resume
                resume_text = extract_text_from_pdf(file)
                if not resume_text:
                    flash("No text extracted from the resume PDF", 'error')
                    return redirect(request.url)

                refined_resume, job_category = refine_resume(job_desc, resume_text)
                
                if refined_resume:
                    # Save the refined resume
                    refined_resume_path = os.path.join(REFINED_DIR, f"refined_resume.docx")
                    docx_buffer = create_docx(refined_resume)
                    
                    if not docx_buffer:
                        flash("Error creating DOCX from refined resume", 'error')
                        return redirect(request.url)

                    with open(refined_resume_path, "wb") as f:
                        f.write(docx_buffer.getvalue())
                    
                    # Save to database and get the resume_id
                    resume_id = save_refined_resume(
                        session['username'],
                        file.filename,
                        resume_text,
                        job_desc,
                        refined_resume_path,
                        job_category
                    )
                    
                    flash('Resume refined successfully!', 'success')
                    return redirect(url_for('feedback', resume_id=resume_id))

            except Exception as e:
                flash(f'Error processing resume: {e}', 'error')
                
    return render_template('new_refinement.html')




@app.route('/my-resumes')
def my_resumes():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT r.*, u.username 
            FROM resumes r
            JOIN users u ON r.user_id = u.id
            WHERE u.username = %s
            ORDER BY r.created_at DESC
        """, (session['username'],))
        resumes = cursor.fetchall()
        
        return render_template('my_resumes.html', resumes=resumes)
    except Error as e:
        flash(f'Database error: {e}', 'error')
        return render_template('my_resumes.html', resumes=[])
    finally:
        if conn:
            conn.close()

# Additional routes for app.py

@app.route('/download-resume/<int:resume_id>')
def download_resume(resume_id):
    if 'username' not in session:
        return redirect(url_for('login'))
        
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Fetch the refined resume path and additional details
        cursor.execute("""
            SELECT r.refined_resume_path, u.username
            FROM resumes r
            JOIN users u ON r.user_id = u.id
            WHERE r.id = %s AND u.username = %s
        """, (resume_id, session['username']))
        resume = cursor.fetchone()
        
        if resume and os.path.exists(resume['refined_resume_path']):
            # Generate a meaningful file name
            username = resume['username'].replace(" ", "_").lower()
            file_name = f"{username}_refined.docx"
            
            return send_file(
                resume['refined_resume_path'],
                as_attachment=True,
                download_name=file_name
            )
        else:
            flash('Resume not found', 'error')
            return redirect(url_for('my_resumes'))
    except Error as e:
        flash(f'Database error: {e}', 'error')
        return redirect(url_for('my_resumes'))
    finally:
        if conn:
            conn.close()
            

import traceback

@app.route('/delete-resume/<int:resume_id>', methods=['POST'])
def delete_resume(resume_id):
    if 'username' not in session:
        return {'success': False, 'error': 'Not authenticated'}, 401

    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        # Check if the resume exists and fetch its file path
        cursor.execute("""
            SELECT id, refined_resume_path 
            FROM resumes 
            WHERE id = %s AND user_id = (SELECT id FROM users WHERE username = %s)
        """, (resume_id, session['username']))
        resume = cursor.fetchone()

        if not resume:
            return {'success': False, 'error': 'Resume not found or not associated with this user'}, 404

        # Log the fetched resume data for debugging
        print(f"Fetched resume: {resume}")
        print(f"Refined resume path: {resume['refined_resume_path']}")

        # Convert the relative file path to an absolute path
        refined_resume_path = os.path.abspath(resume['refined_resume_path'])
        print(f"Absolute file path: {refined_resume_path}")

        # Delete the file if it exists
        if os.path.exists(refined_resume_path):
            try:
                os.remove(refined_resume_path)
                print(f"File deleted: {refined_resume_path}")
            except Exception as e:
                print(f"Failed to delete file {refined_resume_path}: {str(e)}")
                return {'success': False, 'error': f"Failed to delete file: {str(e)}"}, 500
        else:
            print(f"File not found: {refined_resume_path}")

        # Delete any dependent feedback records
        cursor.execute("""
            DELETE FROM feedback WHERE resume_id = %s
        """, (resume_id,))
        
        # Delete the resume row from the database
        cursor.execute("""
            DELETE FROM resumes 
            WHERE id = %s AND user_id = (SELECT id FROM users WHERE username = %s)
        """, (resume_id, session['username']))
        conn.commit()

        # Verify if the resume was deleted from the database
        cursor.execute("""
            SELECT COUNT(*) FROM resumes WHERE id = %s
        """, (resume_id,))
        count = cursor.fetchone()

        if count['COUNT(*)'] == 0:
            print("Resume deleted successfully from the database.")
        else:
            print("Failed to delete resume from the database.")

        return {'success': True}

    except Exception as e:
        # Print the full stack trace for debugging
        error_message = str(e)
        stack_trace = traceback.format_exc()
        print(f"Error occurred: {error_message}")
        print(f"Stack trace: {stack_trace}")
        return {'success': False, 'error': error_message}, 500

    finally:
        if conn:
            conn.close()



@app.route('/submit-feedback', methods=['POST'])
def submit_feedback():
    print(request.form)

    if 'username' not in session:
        flash('Please log in to submit feedback.', 'error')
        return redirect(url_for('login'))
    
    resume_id = request.form.get('resume_id')
    feedback_text = request.form.get('feedback')
    rating = request.form.get('rating')

    if not feedback_text or not rating or not resume_id:
        flash('All fields are required.', 'error')
        return redirect(url_for('my_resumes'))
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Fetch job category from the resume table
        cursor.execute("SELECT job_category FROM resumes WHERE id = %s", (resume_id,))
        job_category = cursor.fetchone()

        if job_category:
            job_category = job_category[0]
        else:
            flash('Resume not found.', 'error')
            return redirect(url_for('my_resumes'))

        cursor.execute("""
            INSERT INTO feedback (resume_id, user_id, feedback_text, rating, created_at)
            VALUES (%s, (SELECT id FROM users WHERE username = %s), %s, %s, NOW())
        """, (resume_id, session['username'], feedback_text, rating))
        conn.commit()

        log_activity(session['username'], f'Submitted feedback for Job Category {job_category}')
        flash('Thank you for your feedback!', 'success')
    except Error as e:
        flash(f'Error saving feedback: {e}', 'error')
    finally:
        if conn:
            conn.close()
            
    return redirect(url_for('my_resumes'))



@app.route('/feedback/<int:resume_id>', methods=['GET'])
def feedback(resume_id):
    if 'username' not in session:
        flash('Please log in to provide feedback.', 'error')
        return redirect(url_for('login'))

    return render_template('feedback.html', resume_id=resume_id)


# Save feedback to CSV
def save_feedback_to_csv(user_id, resume_id, feedback_text, rating):
    feedback_file = "feedback.csv"
    file_exists = os.path.isfile(feedback_file)

    try:
        with open(feedback_file, mode='a', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            if not file_exists:
                writer.writerow(["user_id", "resume_id", "Feedback", "Rating", "Timestamp"])
            writer.writerow([user_id, resume_id, feedback_text, rating, datetime.now()])
    except Exception as e:
        flash(f"Error saving feedback to CSV: {e}", 'error')


def log_activity(username, activity_type, activity_description=None):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Fetch user_id from the username
        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        if user is None:
            flash(f"User {username} not found.", 'error')
            return

        user_id = user[0]  # Extract user ID

        # Insert the activity log
        cursor.execute(
            """
            INSERT INTO activity_logs (user_id, activity_type, activity_description, timestamp)
            VALUES (%s, %s, %s, NOW())
            """,
            (user_id, activity_type, activity_description)
        )
        conn.commit()
    except Error as e:
        print(f"Error logging activity: {e}")
        flash(f"Error logging activity: {str(e)}", 'error')
        if conn:
            conn.rollback()  # Ensure rollback on error
    finally:
        if conn:
            conn.close()



def log_processing_activity(user_id, resume_id, processing_status, processing_time, status):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            """
            INSERT INTO processing_logs 
            (user_id, resume_id, processing_status, processing_time, status, timestamp) 
            VALUES (%s, %s, %s, %s, %s, NOW())
            """,
            (user_id, resume_id, processing_status, processing_time, status)
        )
        conn.commit()
        logger.info(f"Logged processing activity for User ID {user_id}, Resume ID {resume_id}")
    except Error as e:
        logger.error(f"Error logging processing activity: {e}")
        flash(f"Error logging activity: {str(e)}", 'error')
        conn.rollback()  # Ensure rollback on error
    finally:
        if conn:
            conn.close()

            

# Error handler routes for app.py

@app.errorhandler(404)
def not_found_error(error):
    
    app.logger.error(f'Server Error: {error}')
    """Handle 404 Not Found errors"""
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    """Handle 500 Internal Server errors"""
    # Log the error for debugging
    app.logger.error(f'Server Error: {error}')
    # Rollback any failed database sessions

    return render_template('errors/500.html'), 500



if __name__ == '__main__':
    app.run(debug=True)
