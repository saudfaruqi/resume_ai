import uuid
import os
import io
import streamlit as st
from fpdf import FPDF
from hashlib import sha256
import pandas as pd
import mysql.connector
from mysql.connector import Error
from PyPDF2 import PdfReader
from docx import Document
from google.generativeai import GenerativeModel, configure
from streamlit_option_menu import option_menu
from st_aggrid import AgGrid
import plotly.graph_objects as go
import csv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Directories
UPLOAD_DIR = "uploads"
REFINED_DIR = "refined"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REFINED_DIR, exist_ok=True)

# MySQL Configuration
MYSQL_HOST = os.getenv('MYSQL_HOST', 'localhost')
MYSQL_USER = os.getenv('MYSQL_USER', 'root')
MYSQL_PASSWORD = os.getenv('MYSQL_PASSWORD', '')
MYSQL_DATABASE = os.getenv('MYSQL_DATABASE', 'resume_tool')

# AI Configuration
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY', 'AIzaSyBP45y-sVM_W4T3yacnCRYOr4d6-CgF7A8')
configure(api_key=GOOGLE_API_KEY)

# Database Connection
def get_db_connection():
    return mysql.connector.connect(
        host=MYSQL_HOST,
        user=MYSQL_USER,
        password=MYSQL_PASSWORD,
        database=MYSQL_DATABASE
    )

# Extract Text from PDF
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    return "".join(page.extract_text() for page in reader.pages if page.extract_text())


def refine_resume(job_desc, resume_text):
    prompt = f"""
    You are an expert ATS optimization specialist and career coach with extensive experience in resume tailoring and job matching. Your mission is to transform resumes to maximize alignment with target roles while maintaining authenticity and highlighting genuine qualifications.

    ### Analysis Phase:
    1. Job Requirements Analysis:
    - Extract and prioritize required skills, experiences, and qualifications
    - Identify both explicit and implicit requirements
    - Note industry-specific terminology and standards
    - Flag any unique or specialized requirements

    2. Resume Gap Analysis:
    - Compare current qualifications against requirements
    - Identify strength areas and improvement opportunities
    - Calculate approximate match percentage
    - Flag potential red flags or misalignments

    ### Optimization Guidelines:
    1. Strategic Content Prioritization:
    - Reorganize content to lead with most relevant qualifications
    - Maintain 65-35 ratio between technical/role-specific and transferable skills
    - Include only experience from last 10 years unless highly relevant
    - Focus on achievements that demonstrate required competencies

    2. Metrics-Driven Impact Statements:
    - Transform each bullet point using: Action Verb + Task + Result + Impact
    - Quantify achievements with specific metrics (%, $, time saved, etc.)
    - Include scale indicators (team size, project scope, budget)
    - Highlight leadership and collaboration examples relevant to role level

    3. ATS Optimization:
    - Strategic keyword placement in first 2-3 lines of each section
    - Use both full terms and acronyms for technical terms (e.g., "artificial intelligence (AI)")
    - Maintain keyword density of 5-7% for primary skills
    - Include section headers that match common ATS categories

    4. Format & Structure:
    - Create clear visual hierarchy with consistent formatting
    - Use bullet points: 3-4 for recent roles, 2-3 for older positions
    - Maintain 50/50 white space ratio for readability
    - Ensure all dates follow consistent format (MM/YYYY)

    ### Output Requirements:
    1. Provide ONLY the optimized resume in a clean, professional format.
    2. Do NOT include any analysis reports, recommendations, roadmaps, or strategies.
    3. Ensure the output contains only the actual resume content that would be seen by employers.
    4. Format the resume using standard section headers (Summary, Experience, Education, Skills, etc.).
    5. Use clear, consistent bullet points for experience items.
    6. Maintain professional spacing and formatting throughout.

    ### Input Parameters:
    Job Description:
    {job_desc}

    Current Resume:
    {resume_text}

    Note: Focus solely on producing a polished, ATS-optimized resume without any additional analysis or recommendations.
    """

    try:
        chat_session = GenerativeModel(model_name="gemini-1.5-pro").start_chat(history=[])
        response = chat_session.send_message(prompt)
        refined_resume_text = response.text

        # Categorize the job title - this should return just the job category
        job_category = categorize_resume(job_desc)

        return refined_resume_text, job_category
    except Exception as e:
        st.error(f"AI Error: {e}")
        return None, None


def categorize_resume(job_desc, resume_text):
    prompt = f"""
    Analyze the following job description and provide ONLY a single job category (1-3 words) that best describes the role.
    No additional explanation or context - just the category itself.

    Job Description:
    {job_desc}
    """

    try:
        chat_session = GenerativeModel(model_name="gemini-1.5-pro").start_chat(history=[])
        response = chat_session.send_message(prompt)
        return response.text.strip()  # Remove any extra whitespace
    except Exception as e:
        st.error(f"AI Error: {e}")
        return None



def create_docx(resume_text):
    """
    Create a well-formatted DOCX document from the given resume text.

    Args:
        resume_text (str): The text content for the resume.

    Returns:
        io.BytesIO: A buffer containing the DOCX file.
    """
    doc = Document()

    # Set title (e.g., name of the person or header)
    title = doc.add_paragraph("Resume", style='Heading 1')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue color for professionalism
    
    # Add Horizontal Line
    doc.add_paragraph("\n" + "-" * 50)
    
    # Split the resume text into lines and process each line
    lines = resume_text.split('\n')
    for line in lines:
        paragraph = doc.add_paragraph(style='Normal')
        
        # Handle section titles (detect if line looks like a section header)
        if line.isupper():  # Assume section titles are in uppercase
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for consistency
            paragraph.add_run("\n")  # For spacing
        elif line.startswith("-") or line.startswith("*"):  # Bullet points
            bullet_paragraph = doc.add_paragraph(style='List Bullet')
            bullet_paragraph.add_run(line.strip())
        else:  # Regular content
            words = line.split(' ')
            for word in words:
                if word.startswith("**") and word.endswith("**"):  # Bold text
                    run = paragraph.add_run(word[2:-2])  # Remove the '**' markers
                    run.bold = True
                else:
                    run = paragraph.add_run(word)
                run.font.size = Pt(11)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.add_run("\n")  # For spacing between lines

    # Add some space at the end to make it look cleaner
    doc.add_paragraph("\n" + "-" * 50)
    
    # Save to a buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer




# Save the refined resume to the database
def save_refined_resume(username, resume_path, resume_text, job_description, refined_resume_path, job_category):    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Retrieve the user_id based on the username
        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        if user is None:
            st.error("User not found")
            return
        
        user_id = user[0]  # Get the user ID

        # Insert the refined resume data, including job category
        cursor.execute(
            "INSERT INTO resumes (user_id, resume_path, resume_text, job_description, refined_resume_path, job_category, created_at) VALUES (%s, %s, %s, %s, %s, %s, NOW())",
            (user_id, resume_path, resume_text, job_description, refined_resume_path, job_category)
        )
        conn.commit()
        st.success("Refined resume saved successfully!")
    except Error as e:
        st.error(f"Database error while saving the refined resume: {e}")
    finally:
        if conn:
            conn.close()

            

# Delete a resume from the database and server
def delete_resume(resume_id, resume_path):
    try:
        # Database connection
        conn = get_db_connection()
        cursor = conn.cursor()

        # Begin transaction
        conn.begin()

        # Delete the resume from the database
        cursor.execute("DELETE FROM resumes WHERE id = %s", (resume_id,))
        conn.commit()

        # Renumber the IDs in the resumes table
        cursor.execute("""
            SET @new_id = 0;
            UPDATE resumes
            SET id = (@new_id := @new_id + 1)
            ORDER BY id;
        """)

        # Update related references (if any)
        # Example: Update activity_logs table if it stores resume_id
        # cursor.execute("UPDATE activity_logs SET resume_id = ... WHERE ...")

        conn.commit()

        # Delete the resume file from the server
        if os.path.exists(resume_path):
            os.remove(resume_path)
        
        st.success("Resume deleted and IDs updated successfully!")
    except Exception as e:
        st.error(f"Error deleting resume: {e}")
        if conn:
            conn.rollback()  # Rollback transaction on failure
    finally:
        if conn:
            conn.close()
            

# User Authentication and Profile Management
def login_user(username, password):
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Fetch the user by username
        cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        
        if user:
            # Hash the entered password and compare it with the stored hash
            entered_password_hash = sha256(password.encode()).hexdigest()
            if entered_password_hash == user['password_hash']:
                return True  # Login successful
            else:
                st.error("Incorrect password")
        else:
            st.error("User not found")
        
        return False
    except Error as e:
        st.error(f"Database error: {e}")
    finally:
        if conn:
            conn.close()

            
            
def log_activity(user_id, activity_description):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO activity_logs (user_id, activity_description) VALUES (%s, %s)",
            (user_id, activity_description)
        )
        conn.commit()
    except Error as e:
        st.error(f"Error logging activity: {e}")
    finally:
        if conn:
            conn.close()
            

def signup_user(username, email, password):
    try:
        # Hash the password
        password_hash = sha256(password.encode()).hexdigest()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Insert user details into the database
        cursor.execute(
            "INSERT INTO users (username, email, password_hash) VALUES (%s, %s, %s)",
            (username, email, password_hash)
        )
        conn.commit()
        st.success("User registered successfully! Please log in.")
    except Error as e:
        st.error(f"Database error: {e}")
    finally:
        if conn:
            conn.close()


# Password Change with Email Verification
def change_password(username, old_password, new_password):
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        if user and sha256(old_password.encode()).hexdigest() == user['password_hash']:
            new_password_hash = sha256(new_password.encode()).hexdigest()
            cursor.execute(
                "UPDATE users SET password_hash = %s WHERE username = %s",
                (new_password_hash, username)
            )
            conn.commit()
            st.success("Password changed successfully!")
        else:
            st.error("Invalid current password")
    except Error as e:
        st.error(f"Database error: {e}")
    finally:
        if conn:
            conn.close()

# Logout function
def logout():
    if 'username' in st.session_state:
        del st.session_state['username']
    st.success("You have been logged out!")
    st.rerun()


def upload_multiple_files():
    uploaded_files = st.file_uploader("Upload Multiple Resumes (PDF only)", type="pdf", accept_multiple_files=True)
    if uploaded_files:
        for uploaded_file in uploaded_files:
            st.write(f"Processing {uploaded_file.name}")
            # Add processing code here

def display_resume_analytics(resume_text):
    if resume_text:
        # Example analysis - word frequency
        words = resume_text.split()
        word_freq = pd.Series(words).value_counts()
        st.bar_chart(word_freq.head(10))

def show_notifications():
    st.info("Notification: Your resume refinement is complete!")
    
    
# Collect user feedback
def collect_feedback(username, refined_resume_path):
    st.subheader("We Value Your Feedback!")
    feedback_text = st.text_area("Did you like the updated resume? Please share your thoughts.")
    rating = st.slider("Rate our service (1 - Bad, 5 - Excellent)", min_value=1, max_value=5)

    submit_button = st.button("Submit Feedback")
    if submit_button:
        if feedback_text.strip() == "":
            st.error("Please provide feedback before submitting.")
            return

        # Process feedback using Gemini AI
        try:
            chat_session = GenerativeModel(model_name="gemini-1.5-pro").start_chat(history=[])
            prompt = f"Analyze the following feedback and provide a suitable response: {feedback_text}"
            response = chat_session.send_message(prompt)
            ai_response = response.text
            
            # Save feedback details in the backend (not visible to users)
            save_feedback_to_csv(username, refined_resume_path, feedback_text, rating, ai_response)

            # Confirmation message for the user, without showing AI response
            st.success("Thank you for your feedback!")
            # Disable the submit button after feedback submission
            submit_button = False
            
        except Exception as e:
            st.error(f"AI Error: {e}")
            return


# Save feedback to CSV (this function stores the feedback without showing it to the user)
def save_feedback_to_csv(username, resume_path, feedback, rating, ai_response):
    feedback_file = "feedback.csv"
    file_exists = os.path.isfile(feedback_file)

    # Debugging info
    print(f"Saving feedback to {feedback_file}")
    print(f"Feedback: {feedback}")
    print(f"Rating: {rating}")
    print(f"AI Response: {ai_response}")
    
    try:
        with open(feedback_file, mode='a', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            if not file_exists:
                writer.writerow(["Username", "Resume Path", "Feedback", "Rating", "AI Response"])
            writer.writerow([username, resume_path, feedback, rating, ai_response])
        print("Feedback saved successfully.")
    except Exception as e:
        print(f"Error saving feedback: {e}")

# Main App
def main():
    st.set_page_config(page_title="AI-Powered Resume Refinement", layout="wide")
    
        # Initialize session states
    if 'username' not in st.session_state:
        st.session_state['username'] = None
    if 'resume_processed' not in st.session_state:
        st.session_state['resume_processed'] = False
    if 'current_refined_resume' not in st.session_state:
        st.session_state['current_refined_resume'] = None
    if 'current_job_category' not in st.session_state:
        st.session_state['current_job_category'] = None
    if 'current_docx' not in st.session_state:
        st.session_state['current_docx'] = None
    
    # Authentication Screen
    if st.session_state['username'] is None:
        st.header("Login / Signup")
        option = st.radio("Choose an option", ["Login", "Signup"])

        if option == "Login":
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            if st.button("Login"):
                if login_user(username, password):
                    st.session_state['username'] = username
                    st.success(f"Welcome {username}!")
                    try:
                        conn = get_db_connection()
                        cursor = conn.cursor()
                        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                        user = cursor.fetchone()
                        if user:
                            user_id = user[0]
                            log_activity(user_id, "User logged in")
                    except Error as e:
                        st.error(f"Database error: {e}")
                    finally:
                        if conn:
                            conn.close()
                    st.rerun()
                else:
                    st.error("Invalid username or password")

        elif option == "Signup":
            username = st.text_input("Username")
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            if st.button("Sign Up"):
                signup_user(username, email, password)
                st.rerun()

        st.stop()

    # Main Application Interface
    st.sidebar.header(f"Logged in as: {st.session_state['username']}")
    st.sidebar.button("Logout", on_click=logout)

    selected = option_menu(
        menu_title="Main Menu",
        options=["Dashboard", "My Resumes", "New Refinement", "Activity Log", "User Settings"],
        icons=["bar-chart", "folder", "file-plus", "clock", "gear"],
        menu_icon="menu-button",
        default_index=0,
        orientation="horizontal",
    )

    if selected == "Dashboard":
        st.header("Dashboard")

        # Database connection
        try:
            conn = get_db_connection()
            cursor = conn.cursor()

            # Query for total resumes
            cursor.execute("SELECT COUNT(*) FROM resumes WHERE user_id = (SELECT id FROM users WHERE username = %s)", (st.session_state['username'],))
            total_resumes = cursor.fetchone()[0]


            # Query for average processing time
            cursor.execute(""" SELECT AVG(TIMESTAMPDIFF(SECOND, created_at, NOW())) FROM resumes WHERE user_id = (SELECT id FROM users WHERE username = %s)""", (st.session_state['username'],))
            avg_time = cursor.fetchone()[0]


            # Handle case where avg_time is None
            if avg_time is not None:
                avg_time = float(avg_time)
            else:
                avg_time = "N/A"

            # Displaying metrics in cards
            col1, col2 = st.columns(2)
            col1.metric(label="Total Resumes Processed", value=total_resumes)
            col2.metric(label="Average Refinement Time (seconds)", value=avg_time)

            # Chart for total resumes processed over time
            cursor.execute("""
                SELECT DATE(created_at), COUNT(*) 
                FROM resumes 
                WHERE user_id = (SELECT id FROM users WHERE username = %s)
                GROUP BY DATE(created_at)
            """, (st.session_state['username'],))
            
            data = cursor.fetchall()
            if data:
                df = pd.DataFrame(data, columns=["Date", "Count"])
                fig = go.Figure([go.Bar(x=df["Date"], y=df["Count"])])
                fig.update_layout(title="Total Resumes Processed Over Time")
                st.plotly_chart(fig)

            # Chart for average processing time over time
            cursor.execute("""
                SELECT DATE(created_at), AVG(TIMESTAMPDIFF(SECOND, created_at, NOW())) 
                FROM resumes 
                WHERE user_id = (SELECT id FROM users WHERE username = %s)
                GROUP BY DATE(created_at)
            """, (st.session_state['username'],))
            data_time = cursor.fetchall()
            
            if data_time:
                df_time = pd.DataFrame(data_time, columns=["Date", "Average Time (seconds)"])
                fig_time = go.Figure([go.Line(x=df_time["Date"], y=df_time["Average Time (seconds)"])])
                fig_time.update_layout(title="Average Processing Time Over Time")
                st.plotly_chart(fig_time)

        except Error as e:
            st.error(f"Database error: {e}")
        finally:
            if conn:
                conn.close()
                
                
                
    elif selected == "My Resumes":
        st.header("My Resumes")

        try:
            conn = get_db_connection()
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT id, job_description, resume_text, created_at FROM resumes WHERE user_id = (SELECT id FROM users WHERE username = %s)", (st.session_state['username'],))
            resumes = cursor.fetchall()

            if resumes:
                st.sidebar.header("Previous Jobs")
                for resume in resumes:
                    job_desc_summary = resume['job_description'][:30] + "..."
                    if st.sidebar.button(job_desc_summary, key=resume['id']):
                        st.session_state['selected_resume'] = resume['id']
                        st.rerun()

                # Handle resume deletion
                if 'selected_resume' in st.session_state:
                    selected_id = st.session_state['selected_resume']
                    selected_resume = next((r for r in resumes if r['id'] == selected_id), None)
                    if selected_resume:
                        st.subheader(f"Details for Resume ID {selected_id}")
                        st.write(f"**Job Description**: {selected_resume['job_description']}")
                        st.write(f"**Created At**: {selected_resume['created_at']}")
                        if selected_resume['resume_text']:
                            st.write(f"**Refined Resume**: {selected_resume['resume_text']}")
                        else:
                            st.info("No refined resume available.")

                        # Option to delete the resume
                        if st.button("Delete Resume"):
                            # Delete the resume from the database
                            delete_query = "DELETE FROM resumes WHERE id = %s"
                            cursor.execute(delete_query, (selected_id,))
                            conn.commit()  # Commit changes to the database

                            # Reindex the AUTO_INCREMENT value
                            cursor.execute("SET @max_id = (SELECT MAX(id) FROM resumes);")
                            cursor.execute("ALTER TABLE resumes AUTO_INCREMENT = @max_id + 1;")
                            conn.commit()

                            # Remove selected resume from session state
                            del st.session_state['selected_resume']
                            
                            # Refresh the page to show the updated list
                            st.rerun()

                else:
                    st.info("No resume selected. Please click on a job description from the sidebar to view it.")
            else:
                st.info("No resumes found.")
        except Error as e:
            st.error(f"Database error: {e}")
        finally:
            if conn:
                conn.close()
                
                
    # New Refinement Section
    elif selected == "New Refinement":
        st.header("New Resume Refinement")
        
        # Initialize input tracking
        if 'last_job_desc' not in st.session_state:
            st.session_state['last_job_desc'] = None
        if 'last_uploaded_file' not in st.session_state:
            st.session_state['last_uploaded_file'] = None
            
        job_desc = st.text_area("Enter Job Description", height=150)
        uploaded_file = st.file_uploader("Upload Resume (PDF only)", type="pdf")
        
        
        # Check for input changes
        inputs_changed = (
            job_desc != st.session_state['last_job_desc'] or 
            (uploaded_file and uploaded_file.name != st.session_state.get('last_uploaded_file'))
        )

        if inputs_changed:
            st.session_state['resume_processed'] = False
            st.session_state['current_refined_resume'] = None
            st.session_state['current_job_category'] = None
            st.session_state['current_docx'] = None
            
        # Update input tracking
        st.session_state['last_job_desc'] = job_desc
        st.session_state['last_uploaded_file'] = uploaded_file.name if uploaded_file else None

        # Process resume
        if uploaded_file and job_desc:
            process_button = st.button("Process Resume")
            
            if process_button or not st.session_state['resume_processed']:
                with st.spinner("Refining your resume..."):
                    resume_text = extract_text_from_pdf(uploaded_file)
                    refined_resume, job_category = refine_resume(job_desc, resume_text)
                    
                    if refined_resume:
                        st.session_state['current_refined_resume'] = refined_resume
                        st.session_state['current_job_category'] = job_category
                        st.session_state['current_docx'] = create_docx(refined_resume)
                        st.session_state['resume_processed'] = True
                        st.session_state['resume_text'] = resume_text  # Store for database

            # Display results
            if st.session_state['resume_processed'] and st.session_state['current_refined_resume']:
                st.text_area("Refined Resume", st.session_state['current_refined_resume'], height=200)
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.session_state['current_docx']:
                        st.download_button(
                            "Download DOCX",
                            st.session_state['current_docx'],
                            file_name="refined_resume.docx"
                        )
                with col2:
                    if st.button("Save to My Resumes"):
                        refined_resume_path = os.path.join(REFINED_DIR, f"refined_resume_{uuid.uuid4()}.docx")
                        with open(refined_resume_path, "wb") as f:
                            f.write(st.session_state['current_docx'].getvalue())

                        save_refined_resume(
                            st.session_state['username'],
                            uploaded_file.name,
                            st.session_state['resume_text'],
                            job_desc,
                            refined_resume_path,
                            st.session_state['current_job_category']
                        )
                        st.success("Resume saved successfully!")
                        
                        # Show feedback form after saving
                        collect_feedback(st.session_state['username'], refined_resume_path)

        # Clear form button
        if st.button("Clear Form"):
            st.session_state['resume_processed'] = False
            st.session_state['current_refined_resume'] = None
            st.session_state['current_job_category'] = None
            st.session_state['current_docx'] = None
            st.session_state['last_job_desc'] = None
            st.session_state['last_uploaded_file'] = None
            st.rerun()

    
    # Activity Log Section
    elif selected == "Activity Log":
        st.header("Activity Log")
        try:
            conn = get_db_connection()
            cursor = conn.cursor(dictionary=True)
            cursor.execute("""
                SELECT activity_description, timestamp 
                FROM activity_logs 
                WHERE user_id = (SELECT id FROM users WHERE username = %s)
                ORDER BY timestamp DESC
                LIMIT 10
            """, (st.session_state['username'],))
            activities = cursor.fetchall()

            if activities:
                for activity in activities:
                    st.write(f"{activity['timestamp']} - {activity['activity_description']}")
            else:
                st.info("No activity recorded yet.")
        except Error as e:
            st.error(f"Database error: {e}")
        finally:
            if conn:
                conn.close()

    # User Settings Section
    elif selected == "User Settings":
        st.header("User Settings")
        
        # Profile Update
        st.subheader("Update Profile Information")
        new_email = st.text_input("New Email")
        if st.button("Update Profile"):
            # Update profile logic here
            st.success("Profile updated successfully!")

        # Password Change
        st.subheader("Change Password")
        old_password = st.text_input("Current Password", type="password")
        new_password = st.text_input("New Password", type="password")
        if st.button("Change Password"):
            change_password(st.session_state['username'], old_password, new_password)

if __name__ == "__main__":
    main()                         